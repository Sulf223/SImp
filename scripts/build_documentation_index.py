#!/usr/bin/env python3
"""
Build a small searchable text index from proiect_documentatie.

The PHP app runs from Docker with only site_g mounted, so the extracted
documentation context is stored in site_g/storage/documentation_index.json.
"""

from __future__ import annotations

import hashlib
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path


SUPPORTED_EXTENSIONS = {".cpp", ".txt", ".pdf", ".docx"}
MAX_CHARS_PER_FILE = 45000
CHUNK_SIZE = 1800
CHUNK_OVERLAP = 220


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "cp1250", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - local tooling guard
        raise RuntimeError("pypdf is required for PDF extraction") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for idx, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(f"[pagina {idx + 1}]\n{text}")
    return "\n\n".join(pages)


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - local tooling guard
        raise RuntimeError("python-docx is required for DOCX extraction") from exc

    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".cpp", ".txt"}:
        return read_text_file(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    return ""


def make_title(path: Path, root: Path) -> str:
    relative = path.relative_to(root)
    stem = path.stem.replace("_", " ").replace("-", " ")
    parent = relative.parent.name.replace("_", " ").replace("-", " ")
    return f"{stem} ({parent})" if parent and parent != "." else stem


def split_chunks(text: str) -> list[str]:
    if len(text) <= CHUNK_SIZE:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        if end < len(text):
            boundary = max(text.rfind("\n\n", start, end), text.rfind("\n", start, end), text.rfind(". ", start, end))
            if boundary > start + 500:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(0, end - CHUNK_OVERLAP)
    return chunks


def build_index(repo_root: Path) -> dict:
    docs_root = repo_root / "proiect_documentatie"
    if not docs_root.exists():
        raise FileNotFoundError(f"Missing documentation directory: {docs_root}")

    chunks: list[dict] = []
    seen_hashes: set[str] = set()
    errors: list[dict] = []

    for path in sorted(docs_root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        relative = path.relative_to(repo_root).as_posix()
        try:
            text = clean_text(extract_text(path))
        except Exception as exc:
            errors.append({"source": relative, "error": str(exc)})
            continue

        if not text:
            continue

        text = text[:MAX_CHARS_PER_FILE]
        title = make_title(path, docs_root)
        for index, chunk in enumerate(split_chunks(text)):
            fingerprint = hashlib.sha256(chunk.encode("utf-8", errors="ignore")).hexdigest()
            if fingerprint in seen_hashes:
                continue
            seen_hashes.add(fingerprint)
            chunks.append(
                {
                    "source": relative,
                    "title": title,
                    "type": path.suffix.lower().lstrip("."),
                    "chunk": index,
                    "text": chunk,
                }
            )

    return {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "source_root": "proiect_documentatie",
        "chunk_count": len(chunks),
        "chunks": chunks,
        "errors": errors,
    }


def main() -> int:
    repo_root = Path(__file__).resolve().parents[1]
    output = repo_root / "site_g" / "storage" / "documentation_index.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    index = build_index(repo_root)
    output.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {index['chunk_count']} chunks to {output}")
    if index["errors"]:
        print(f"Extraction warnings: {len(index['errors'])}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
